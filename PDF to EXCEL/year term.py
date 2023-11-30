import re
from docx import Document

def extract_words_before_table(doc):
    pattern = r'Bills - (\w+).*?year (\w+)'

    for paragraph in doc.paragraphs:
        for table in doc.tables:
            if paragraph in table._element.iterancestors('w:tbl'):
                return None, None  # Stop searching when a paragraph is part of a table

        match = re.search(pattern, paragraph.text)
        if match:
            word_after_bills = match.group(1)
            word_after_year = match.group(2)
            return word_after_bills, word_after_year

    return None, None

# Replace 'your_document.docx' with the actual path to your DOCX file
doc_path = 'E:\New\Demo.docx'

try:
    doc = Document(doc_path)
    result_bills, result_year = extract_words_before_table(doc)
    
    if result_bills and result_year:
        print("Word after 'bills -':", result_bills)
        print("Word after 'year':", result_year)
    else:
        print("No match found before the table.")
except FileNotFoundError:
    print(f"Error: The file '{doc_path}' was not found.")
except Exception as e:
    print(f"Error: {e}")
