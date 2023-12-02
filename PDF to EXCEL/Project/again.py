import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import pandas as pd
import os
import shutil
import docx
from docx import Document
from openpyxl import load_workbook
from openpyxl import Workbook
import xlwings as xw
import re
from docx import Document
from num2words import num2words
from googletrans import Translator
from functools import partial
import time
import bangla
import threading
from googletrans import Translator
# <---------------------------------------------- import libraries ----------------------------------------------> end


# <---------------------------------------------- class section ----------------------------------------------> start


class WordToExcelConverter:

# <---------------------------------------------- definition section ----------------------------------------------> start

    def __init__(self, root):
        self.root = root
        self.root.title("Automatic bill generator")
        self.docx_file = None
        self.sample_excel = None
        self.output_dir = None
        self.tables_with_titles = None
        self.combined_excel_path = None
        self.file_handling_thread = None
        self.name = ""
        self.dean=""
        self.head=""
        self.online=1
        # self.file_list = []
        self.committee = []
        self.new_files = []  # Array to store new_file values globally
        self.extracted_data = []
        # self.paused = False
        self.total_no_of_table=0
        self.year=0
        self.term=0
        self.AD=0
        self.dept="সিএসই"    # Mother Department
        self.translator = Translator()
        self.dept_suffixes_mapping = {
             
        "computer science and engineering": "সিএসই",
        "computer science & engineering": "সিএসই",
        "electrical and electronic engineering": "ইইই",
        "electrical & electronic engineering": "ইইই",
        "electronics and communication engineering": "ইসিই",
        "electronics & communication engineering": "ইসিই",
        "biomedical engineering": "বিএমই",
        "materials science and engineering": "এমএসই",
        "materials science & engineering": "এমএসই",
        "civil engineering": "পুরকৌশল",
        "urban and regional planning": "ইউআরপি",
        "urban & regional planning": "ইউআরপি",
        "building engineering and construction management": "বিইসিএম",
        "building engineering & construction management": "বিইসিএম",
        "architecture": "স্থাপত্য",
        "mathematics": "গণিত",
        "math": "গণিত",
        "chemistry": "রসায়ন",
        "physics": "পদার্থ",
        "humanities": "মানবিক",
        "mechanical engineering": "যন্ত্র প্রকৌশল",
        "industrial engineering and management": "শিল্প প্রকৌশল",
        "industrial engineering & management": "শিল্প প্রকৌশল",
        "energy science and engineering": "ইএসই",
        "energy science & engineering": "ইএসই",
        "leather engineering": "লেদার",
        "textile engineering": "টেক্সটাইল",
        "chemical engineering": "টেক্সটাইল",
        "mechatronics engineering": "মেকাট্রনিক্স",
        }
        self.setup_gui()

# <---------------------------------------------- GUI section ----------------------------------------------> start
                
    def setup_gui(self):
        # Create main frame
        main_frame = tk.Frame(self.root)
        main_frame.pack()
        self.main_frame= main_frame

        # Create top frame for title
        top_frame = tk.Frame(main_frame, bg='white')
        top_frame.pack(fill=tk.X)
        self.top_frame= top_frame

        # Title label with mixed colors
        title_label = tk.Label(top_frame, text="Automatic bill generator", font=('Arial', 18, 'bold'), bg='white')
        title_label.pack(padx=300, pady=10)
        # Change text color by segments
        title_label.config(fg='#0000FF')  # Blue color

        # Create middle frame for left and right sections
        middle_frame = tk.Frame(main_frame)
        middle_frame.pack(fill=tk.BOTH, expand=True)
        self.middle_frame= middle_frame

        # Left frame for existing content
        left_frame = tk.Frame(middle_frame)
        left_frame.pack(side=tk.LEFT, padx=10, pady=10, fill=tk.BOTH, expand=True)
        self.left_frame= left_frame

        # Existing content - Add your current UI elements here
        select_button = ttk.Button(left_frame, text="Select Word Doc", command=self.select_docx)
        select_button.pack(pady=10)

        select_sample_button = ttk.Button(left_frame, text="Select Sample Excel", command=self.select_sample_excel)
        select_sample_button.pack(pady=10)


        # Entry widget to take text input

        self.entry = tk.Entry(left_frame, width=30)
        self.entry.pack()  # Pack entry widget to the left side as well

        def update_name():
            self.name = self.entry.get()
            print("Name: ", self.name)
            # self.update_label_text()
            self.generate_excel_from_docx(1)

        # def update_label_text():
        #     self.label.config(text=f"Entered Name: {self.name}")

        # Button to update the label text
        self.update_button = tk.Button(left_frame, text="Generate Individuals Bill", command=update_name)
        self.update_button.pack()

        # Label to display the input text
        # self.label = tk.Label(left_frame, text="Enter text in the Entry and click 'Update Label'")
        # self.label.pack()

        generate_button = ttk.Button(left_frame, text="Generate Bill For all Teachers", command=partial(self.generate_excel_from_docx,0))
        generate_button.pack(pady=10)

        # process_button = ttk.Button(left_frame, text="Process the first table", command=self.process_first_table)
        # process_button.pack(pady=10)

        self.file_handling_thread = None
        self.pause_event = threading.Event()

        # Create pause, continue, and reset buttons
        self.pause_button = tk.Button(left_frame, text="Pause", command=self.pause_progress)
        self.continue_button = tk.Button(left_frame, text="Continue", command=self.continue_progress)
        self.reset_button = tk.Button(left_frame, text="Reset", command=self.reset_progress)
        # Pack buttons in a horizontal line
        self.pause_button.pack(side=tk.LEFT, padx=5, pady=10)
        self.continue_button.pack(side=tk.LEFT, padx=5, pady=10)
        self.reset_button.pack(side=tk.LEFT, padx=5, pady=10)
        
        # Pack buttons in a horizontal line
        self.pause_button.pack(side=tk.LEFT, padx=5, pady=10)
        self.continue_button.pack(side=tk.LEFT, padx=5, pady=10)
        self.reset_button.pack(side=tk.LEFT, padx=5, pady=10)
        # Generate other UI elements as needed in the left_frame...

        # Right frame for empty area to be utilized
        right_frame = tk.Frame(middle_frame, bg='lightgray')
        right_frame.pack(side=tk.RIGHT, padx=10, pady=10, fill=tk.BOTH, expand=True)
        self.right_frame= right_frame


        # Store area in the right frame
        # Create a Listbox widget to display the list in the right frame
        self.listbox = tk.Listbox(right_frame)
        self.listbox.pack(padx=10, pady=10, fill=tk.BOTH, expand=True)

        bottom_frame = tk.Frame(main_frame)
        bottom_frame.pack(fill=tk.BOTH, expand=True)
        self.bottom_frame= bottom_frame


        self.docx_label = tk.Label(bottom_frame, text="Selected Word Doc: ")
        self.docx_label.pack()

        self.sample_label = tk.Label(bottom_frame, text="Selected Sample Excel: ")
        self.sample_label.pack()

        self.progress_bar = ttk.Progressbar(bottom_frame, orient=tk.HORIZONTAL, length=400, mode='determinate')
        self.progress_bar.pack()
        self.progress_bar.pack_forget()


# <---------------------------------------------- GUI section ----------------------------------------------> end

    def update_listbox(self):
        self.listbox.delete(0, tk.END)  # Clear the Listbox before updating
        for file in self.new_files:
            self.listbox.insert(tk.END, file)

    def start_file_handling(self):
        self.root.after(100, self.process_first_table)
         # self.file_handling_thread = threading.Thread(target=self.process_first_table)
        # self.file_handling_thread.start()

    def pause_progress(self):
        self.pause_event.set()
        self.pause_button.config(state=tk.DISABLED)
        self.continue_button.config(state=tk.NORMAL)

    def continue_progress(self):
        self.pause_event.clear()
        self.continue_button.config(state=tk.DISABLED)
        self.pause_button.config(state=tk.NORMAL)
        if self.file_handling_thread and not self.file_handling_thread.is_alive():
            self.start_file_handling()

    def reset_progress(self):
        # Reset operation to initial state
        self.pause_event.clear()
        self.pause_button.config(state=tk.NORMAL)
        self.continue_button.config(state=tk.DISABLED)
        if self.file_handling_thread and self.file_handling_thread.is_alive():
            self.file_handling_thread.join()
        # Reset other necessary states or variables       
        self.clear_labels() 

    def update_progress_bar(self, value):
        self.progress_bar['value'] = value
        self.root.update_idletasks()  # Refresh the window to update the progress bar
    
    def update_docx_label(self):
        if self.docx_file:
            self.docx_label.config(text=f"Selected Word Doc: {self.docx_file}")
            self.docx_label.pack()

    def update_sample_label(self):
        if self.sample_excel:
            self.sample_label.config(text=f"Selected Sample Excel: {self.sample_excel}")
            self.sample_label.pack()

    def display_table_data(self, table_data):
        pass

    def show_error_message(self, message):
        messagebox.showerror("Error", message)

    def show_success_message(self, message):
        messagebox.showinfo("Success", message)

    def pause_execution(self):
        while self.paused:
            time.sleep(1)

    def toggle_pause(self):
        self.paused = not self.paused
        if self.paused:
            self.pause_button.config(state="disabled")
            self.continue_button.config(state="active")
        else:
            self.pause_button.config(state="active")
            self.continue_button.config(state="disabled")

    def clear_labels(self):
        self.docx_label.config(text="Selected Word Doc: ")
        self.docx_label.pack()
        self.sample_label.config(text="Selected Sample Excel: ")
        self.sample_label.pack()
        self.progress_bar.pack_forget()  # Hide the progress bar
        self.update_progress_bar(0)
        self.entry.delete(0, tk.END) 
        self.listbox.delete(0, tk.END)
        self.docx_file = None
        self.sample_excel = None
        self.output_dir = None
        self.tables_with_titles = None
        self.combined_excel_path = None
        self.file_handling_thread = None
        self.name = ""
        self.new_files = [] 
        self.extracted_data = []
        self.total_no_of_table=0
        self.year=0
        self.term=0
        self.AD=0
        self.dept=""

    def english_to_bengali_number_in_words(self, english_number):
         # if offline, cant use google translator api
        if self.online==0:
            return english_number
        try:
            translator = Translator()
            words_in_english = num2words(english_number, lang='en_IN')
            words_in_bengali = translator.translate(words_in_english, dest='bn').text
            modified_output = words_in_bengali.replace(',', '') + " টাকা মাত্র।"
            return modified_output
        except Exception as e:
            # Handle translation errors (including timeouts) here
            print("Translation error:", e)
            return None  # Set a default value or handle the failure accordingly

    def should_skip_translation(self, text):
        name_patterns = [r'Dean', r'Md\.', r'Dr\.', r'Sk\.', r'Fatema']
        for pattern in name_patterns:
            if re.search(pattern, text):
                return True
        return False

    def translate_to_bengali(self, text):
        translator = Translator()
        # if offline, cant use google translator api
        if self.online==0:
            return text
    
        # Define the translation rules
        translation_rules = {
            r'Dean': 'ডিন',
            r'Md\.': 'মোঃ',
            r'Dr\.': 'ড.',
            r'Sk\.': 'শেখ',
            r'Most': 'মোসাম্মৎ',
            r'Fatema': 'ফাতেমা'
        }

        parts = text.split()
        translated_parts = []
        for part in parts:
            if not self.should_skip_translation(part):
                # Apply the specific translation rule if found
                for pattern, replacement in translation_rules.items():
                    if re.search(pattern, part):
                        part = re.sub(pattern, replacement, part)
                        break
                translated_part = translator.translate(part, dest='bn').text
            else:
                # Use provided translation rules when skipping translation
                for pattern, replacement in translation_rules.items():
                    if re.search(pattern, part):
                        translated_part = re.sub(pattern, replacement, part)
                        break
            translated_parts.append(translated_part)

        return ' '.join(translated_parts)

    def print_matching_value_for_file(self, new_file, name, designation, department):
        print("Processing...")
        total_no_of_table = len(self.tables_with_titles) #13
        print("Total no of table: ", total_no_of_table)

        # An array of the size of the total_no_of_table, initially all value is 0
        matching_values = [0] * total_no_of_table #13
        matching_values_2D = [[] for _ in range(total_no_of_table)]



        # Set Name, Year, Term
        name=name.split('(')[0]
        print("Name: ",name)
        # name = self.translate_to_bengali(name)
        # print("Name: ",name)

        print("Designation: ",designation)
        designation = self.translate_to_bengali(designation)
        print("Designation: ",designation)

        print("Department: ",department)
        department = self.dept_translate_to_bengali(department.lower())
        print("Department: ",department)




        # Question Paper Setter & Script Examiner 
        if total_no_of_table > 1:
            table_data = self.tables_with_titles[1]["Table"]
            table_df = pd.DataFrame(table_data)
            # print("Lets see: ")
            matching_values_2D[1].extend([0]*4)
            count_g9=0
            for row_idx in range(1, len(table_df)):
                table_value = str(table_df.iloc[row_idx, 1]).replace(" ", "").replace(".", "").replace(",", "")
                # print(table_value)
                if str(new_file).lower() in table_value.lower() or table_value.lower() in str(new_file).lower():
                    count_g9+=1
                    matching_values_2D[1][3] += float(table_df.iloc[row_idx, 3]) 
                    print(f"Question Paper Setter & Script Examiner : Matching value for {new_file}: {matching_values_2D[1][3]}")

            matching_values_2D[1][0]=count_g9/2
            print(f"Question Paper Setter & Script Examiner : Matching value for {new_file}: {matching_values_2D[1][0]}")

            # Question Moderation Committee
            if name in self.head or self.head in name:
                matching_values_2D[1][1]=1
                print("Question Moderation Committee: Committe Chairman")
            if name in self.committee:
                print("Question Moderation Committee: In Committee")
                matching_values_2D[1][2]=1
                

        # Examiners of Class Tests
        if total_no_of_table > 2:
            table_data = self.tables_with_titles[2]["Table"]
            table_df = pd.DataFrame(table_data)

            matching_values_2D[2].extend([0]*2)
            for row_idx in range(1, len(table_df)):
                table_value = str(table_df.iloc[row_idx, 1]).replace(" ", "").replace(".", "").replace(",", "")
                if str(new_file).lower() in table_value.lower() or table_value.lower() in str(new_file).lower():
                    matching_values_2D[2][0] += float(table_df.iloc[row_idx, 2]) 
                    print(f"Examiners of Class Tests: Matching value for {new_file}: {matching_values_2D[2][0]}")

            matching_values_2D[2][1]=1.5
            print(f"Examiners of Class Tests: Matching value for {new_file}: {matching_values_2D[2][1]}")

        # Examiners of Sessional Classes
        if total_no_of_table > 3:
            table_data = self.tables_with_titles[3]["Table"]
            table_df = pd.DataFrame(table_data)
            
            matching_values_2D[3].extend([0]*2)
            for row_idx in range(1, len(table_df)):
                table_value = str(table_df.iloc[row_idx, 1]).replace(" ", "").replace(".", "").replace(",", "")
                if str(new_file).lower() in table_value.lower() or table_value.lower() in str(new_file).lower():
                    matching_values_2D[3][0] += float(table_df.iloc[row_idx, 2])
                    matching_values_2D[3][1] += float(table_df.iloc[row_idx, 3])/1.5
                    print(f"Examiners of Sessional Classes: Matching value for {new_file}: {matching_values_2D[3][0]} & {matching_values_2D[3][1]}")

        # Script Scrutinizer
        if total_no_of_table > 4:
            table_data = self.tables_with_titles[4]["Table"]
            table_df = pd.DataFrame(table_data)

            matching_values_2D[4].extend([0]*1)
            for row_idx in range(1, len(table_df)):
                table_value = str(table_df.iloc[row_idx, 0]).replace(" ", "").replace(".", "").replace(",", "")
                if str(new_file).lower() in table_value.lower() or table_value.lower() in str(new_file).lower():
                    matching_values_2D[4][0] += float(table_df.iloc[row_idx, 1]) 
                    print(f"Script Scrutinizer: Matching value for {new_file}: {matching_values_2D[4][0]}")


        # Tabulation & Verification
        if total_no_of_table > 5:
            table_data = self.tables_with_titles[5]["Table"]
            table_df = pd.DataFrame(table_data)

            matching_values_2D[5].extend([0]*2)
            for row_idx in range(1, len(table_df)):
                table_value = str(table_df.iloc[row_idx, 1]).replace(" ", "").replace(".", "").replace(",", "")
                if str(new_file).lower() in table_value.lower() or table_value.lower() in str(new_file).lower():
                    matching_values_2D[5][0] += float(table_df.iloc[row_idx, 2]) 
                    matching_values_2D[5][1] += float(table_df.iloc[row_idx, 2]) 
                    print(f"Tabulation & Verification: Matching value for {new_file}: {matching_values_2D[5][0]} & {matching_values_2D[5][0]}")


        # Typing and Drawing
        if total_no_of_table > 6:
            table_data = self.tables_with_titles[6]["Table"]
            table_df = pd.DataFrame(table_data)

            matching_values_2D[6].extend([0]*1)
            for row_idx in range(1, len(table_df)):
                table_value = str(table_df.iloc[row_idx, 0]).replace(" ", "").replace(".", "").replace(",", "")
                if str(new_file).lower() in table_value.lower() or table_value.lower() in str(new_file).lower():
                    matching_values_2D[6][0] += float(table_df.iloc[row_idx, 1]) 
                    print(f"Typing and Drawing: Matching value for {new_file}: {matching_values_2D[6][0]}")


        # Central Viva-Voce
        if total_no_of_table > 7:
            table_data = self.tables_with_titles[7]["Table"]
            table_df = pd.DataFrame(table_data)
            count_h18=0
            matching_values_2D[7].extend([0]*2)
            for row_idx in range(1, len(table_df)):
                count_h18+=1
                table_value = str(table_df.iloc[row_idx, 1]).replace(" ", "").replace(".", "").replace(",", "")
                if str(new_file).lower() in table_value.lower() or table_value.lower() in str(new_file).lower():
                    matching_values_2D[7][0] += float(table_df.iloc[row_idx, 2]) 
                    print(f"Central Viva-Voce: Matching value for {new_file}: {matching_values_2D[7][0]}")
             
            matching_values_2D[7][1] = count_h18


        # Student Advising
        if total_no_of_table > 8:
            table_data = self.tables_with_titles[8]["Table"]
            table_df = pd.DataFrame(table_data)

            matching_values_2D[8].extend([0]*1)
            for row_idx in range(1, len(table_df)):
                table_value = str(table_df.iloc[row_idx, 1]).replace(" ", "").replace(".", "").replace(",", "")
                if str(new_file).lower() in table_value.lower() or table_value.lower() in str(new_file).lower():
                    matching_values_2D[8][0] += float(table_df.iloc[row_idx, 2]) 
                    print(f"Student Advising: Matching value for {new_file}: {matching_values_2D[8][0]}")


        # Seminar (CSE 4120) 1 + 1 =2
        if total_no_of_table > 9:
            table_data = self.tables_with_titles[9]["Table"]
            table_df = pd.DataFrame(table_data)

            matching_values_2D[9].extend([0]*2)
            for row_idx in range(1, len(table_df)):
                table_value = str(table_df.iloc[row_idx, 1]).replace(" ", "").replace(".", "").replace(",", "")
                if str(new_file).lower() in table_value.lower() or table_value.lower() in str(new_file).lower():
                    matching_values_2D[9][0] += float(table_df.iloc[row_idx, 2]) 
                    print(f"Seminar (CSE 4120): Matching value for {new_file}: {matching_values_2D[9][0]}")
                    no_of_student=int(table_df.iloc[row_idx, 2])
                    count_h16=0
                    temp=0
                    for row_idx2 in range(1, len(table_df)):
                        if no_of_student==int(table_df.iloc[row_idx2, 2]):
                            temp += 1
                            count_h16=temp
                        elif row_idx>row_idx2:
                            temp=0
                            count_h16=0
                        else:
                            break
                    matching_values_2D[9][1] +=count_h16
                    print(f"Seminar (CSE 4120): Matching value for {new_file}: {matching_values_2D[9][1]}")


        # Thesis Progress Defense
        if total_no_of_table > 10:
            table_data = self.tables_with_titles[10]["Table"]
            table_df = pd.DataFrame(table_data)

            matching_values_2D[10].extend([0]*2)
            for row_idx in range(1, len(table_df)):
                table_value = str(table_df.iloc[row_idx, 1]).replace(" ", "").replace(".", "").replace(",", "")
                if str(new_file).lower() in table_value.lower() or table_value.lower() in str(new_file).lower():
                    matching_values_2D[10][0] += float(table_df.iloc[row_idx, 2]) 
                    print(f"Thesis Progress Defense: Matching value for {new_file}: {matching_values_2D[10][0]}")
                    no_of_student=int(table_df.iloc[row_idx, 2])
                    count_h20=0
                    temp=0
                    for row_idx2 in range(1, len(table_df)):
                        if no_of_student==int(table_df.iloc[row_idx2, 2]):
                            temp += 1
                            count_h20=temp
                        elif row_idx>row_idx2:
                            temp=0
                            count_h20=0
                        else:
                            break
                    matching_values_2D[10][1] +=count_h20
                    print(f"Thesis Progress Defense: Matching value for {new_file}: {matching_values_2D[10][1]}")


        # Final Grade Sheet Verification
        if total_no_of_table > 11:
            table_data = self.tables_with_titles[11]["Table"]
            table_df = pd.DataFrame(table_data)

            matching_values_2D[11].extend([0]*1)
            for row_idx in range(1, len(table_df)):
                table_value = str(table_df.iloc[row_idx, 1]).replace(" ", "").replace(".", "").replace(",", "")
                if str(new_file).lower() in table_value.lower() or table_value.lower() in str(new_file).lower():
                    print(table_df)
                    matching_values_2D[11][0] += float(table_df.iloc[row_idx, 2]) 
                    print(f"Final Grade Sheet Verification: Matching value for {new_file}: {matching_values_2D[11][0]}")

        
        # Thesis Progress Defense
        if total_no_of_table > 12:
            table_data = self.tables_with_titles[12]["Table"]
            table_df = pd.DataFrame(table_data)

            matching_values_2D[12].extend([0]*2)
            for row_idx in range(1, len(table_df)):
                table_value = str(table_df.iloc[row_idx, 1]).replace(" ", "").replace(".", "").replace(",", "")
                if str(new_file).lower() in table_value.lower() or table_value.lower() in str(new_file).lower():
                    matching_values_2D[12][0] += float(table_df.iloc[row_idx, 2]) 
                    print(f"Thesis Progress Defense: Matching value for {new_file}: {matching_values_2D[12][0]}")
                
            # print(new_file, " ", self.dean.replace(" ", "").replace(".", "").replace(",", ""))
            if new_file in self.dean.replace(" ", "").replace(".", "").replace(",", "") or self.dean.replace(" ", "").replace(".", "").replace(",", "") in new_file:
                 matching_values_2D[12][1] = 3450   # for dean: 3450 taka, for others: given(2700) 
                 print(f"Thesis Progress Defense: Matching value for {new_file}: {matching_values_2D[12][1]}")




        # Determine which bills is needed to write in which cell.
        file_path = os.path.join(self.output_dir, f"{new_file}.xlsx")
        cell_mappings = {
            1: ['G9', 'G10','G11','G12'],
            2: ['G14','H14'],
            3: ['G17','H17'],
            4: ['G25'],
            5: ['G23', 'G24'],
            6: ['G27'],
            7: ['G18','H18'],
            8: ['G29'],
            9: ['G16','H16'],
            10: ['G20','H20'],
            11: ['G28'],
            12: ['G26','K26']
        }

        wb = xw.Book(file_path)
        sheet = wb.sheets.active
        for i in range(1, total_no_of_table):
            if os.path.exists(file_path):
                cell_locations = cell_mappings.get(i, [])
                j=0
                for cell in cell_locations:
                    try:
                        if matching_values_2D[i][j] != 0:
                            print("Inserting data at ", f"{new_file}.xlsx")
                            sheet.range(cell).value = matching_values_2D[i][j]
                    except Exception as e:
                        print(f"Error processing file {file_path}: {e}")
                    j+=1
            else:
                print(f"File {file_path} does not exist.")
        sheet['A3'].value = "নাম: " + name
        sheet['A4'].value = "পদবী: " + designation
        sheet['B5'].value = department
        # wb.save(file_path)
        # wb.close()


        # # Open the workbook
        # wb = xw.Book(file_path)
        # sheet = wb.sheets.active
        # Read value amount from cell I31
        amount_str = sheet.range('I32').value
        english_str = str(amount_str).split('.')[0]    #type casting the float into string and taking the integer portion only
        english_number= int(english_str)
        bengali_words = self.english_to_bengali_number_in_words(english_number)
        sheet['A32'].value = sheet['A32'].value + str(bengali_words)
        wb.save(file_path)
        wb.close()
        print(bengali_words)

        print(matching_values_2D)

    def process_first_table(self):
        self.update_progress_bar(2)
        if self.output_dir and self.tables_with_titles and self.sample_excel:
            file_count = 0  # Counter for the files being created
            combined_df = pd.DataFrame()  # Initialize an empty DataFrame to hold all tables

            for i, data in enumerate(self.tables_with_titles):
                table = data["Table"]
                df = pd.DataFrame(table)
                df.ffill(axis=0, inplace=True)
                if i == 0:  # Working with the first table
                    first_table_df_name = df.iloc[:, 1]  # Extracting the content from the second column
                    first_table_df_designation = df.iloc[:, 2]  # Extracting the content from the second column

                    # Create separate Excel files based on each row's content
                    temp_folder = os.path.join(self.output_dir, "AllTables")
                    if not os.path.exists(temp_folder):
                        os.makedirs(temp_folder)
                    
                    # Create separate Excel files based on each row's content
                    for row_i, (name, designation_and_department) in enumerate(zip(first_table_df_name,first_table_df_designation)):

                        if row_i != 0 and row_i != len(first_table_df_name) - 1:
                            name=name.split(',')[0]
                            designation=designation_and_department.split(',')[0]
                            department=designation_and_department.split(',')[1]
                            if "dean" in designation.lower():
                                self.dean=name
                                print("Dean Name: ",self.dean)
                            print(name, " ", self.name)
                            if self.name.lower() in name.lower() or name.lower() in self.name.lower():
                                new_file=name.replace(" ", "").replace(".", "").replace(",", "")
                                self.new_files.append(new_file)  # Append new_file to the global array
                                new_file_name = new_file + ".xlsx"
                                file_path = os.path.join(self.output_dir, new_file_name)
                                print(f"Creating {new_file_name}... at {file_path}")
                                # self.file_list.append(new_file)
                                shutil.copy(self.sample_excel, file_path)
                                self.print_matching_value_for_file(new_file, name, designation, department)
                                file_count += 1  # Increment file count
                                self.update_progress_bar(file_count*3)

                    # Append the first table content to the combined DataFrame
                    combined_df = pd.concat([combined_df, first_table_df_name], axis=1)
                else:
                    combined_df = pd.concat([combined_df, df.iloc[1:-1, 1]], axis=1)
            self.update_listbox()
            self.update_progress_bar(100)
            print("The total no of files are:", file_count)
            print("The files are:", self.new_files)
            messagebox.showinfo("Congratulations!", f"Excel Created Successfully! Total Files Created: {file_count}")
            self.progress_bar.pack_forget()  # Hide the progress bar
            self.update_progress_bar(0)
            # self.clear_labels()
        else:
            messagebox.showwarning("No Tables Found or No Sample Excel", "No tables were detected in the Word document or no Sample Excel selected.")

    def convert_year_term_suffixes_to_bengali(self, text):
        # Dictionary mapping English year_term_suffixes to Bengali
        year_term_suffixes_mapping = {
            "1st": "১ম",
            "2nd": "২য়",
            "3rd": "৩য়",
            "4th": "৪র্থ",  # You can add more mappings as needed
            # Add more mappings for other year_term_suffixes
        }

        # Replace English year_term_suffixes with Bengali equivalents
        for suffix in year_term_suffixes_mapping:
            if suffix in text:
                text = text.replace(suffix, year_term_suffixes_mapping[suffix])

        return text

    def extract_year_and_term(self):
        pattern_y_t = r'Bills - (\w+).*?year (\w+)'
        pattern_e = r'Examination- (\w+)'
        doc = Document(self.docx_file)
        for paragraph in doc.paragraphs:
            match_bills = re.search(pattern_y_t, paragraph.text)
            match_exam = re.search(pattern_e, paragraph.text)
        
            if match_bills:
                self.year = match_bills.group(1)
                self.term = match_bills.group(2)
                print("Year & Term extracted successfully!")

            if match_exam:
                self.AD = match_exam.group(1)
                print("AD extracted successfully!")

    def dept_translate_to_bengali(self, english_text):
        bengali_text = self.dept_suffixes_mapping.get(english_text.lower())
        if not bengali_text:
            # If the translation is not found in the mapping, use Google Translate
            # if offline
            if self.online==0:
                return english_text
            translated = self.translator.translate(english_text, dest='bn')
            bengali_text = translated.text
        return bengali_text

    def extract_department_line(self):
        pattern = r'(?:Department of|Department Of)(.*)'
        doc = Document(self.docx_file)
        for paragraph in doc.paragraphs:
            match = re.search(pattern, paragraph.text)
            if match:
                self.dept = match.group(1).strip()  # Extract text after the department pattern
                return

        return None

    def extract_information_from_docx(self):
        # Define regular expression pattern for extracting structured information
        pattern = re.compile(r"Dr\..*?, (Professor.*?), (Dept\. of .*?), (KUET.*?)\s+(Ext\.\s+Member|Member|Chairman)$")

        # Load the Word document
        doc = docx.Document(self.docx_file)

        # Initialize lists to store extracted information
        names = []
        titles = []
        departments = []
        institutions = []
        roles = []

        # Iterate through paragraphs in the document and extract information using regex
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            # Apply regex pattern to extract information
            match = pattern.match(text)
            if match:
                names.append(text.split(',')[0].strip())  # Extracting name separately due to different structure
                titles.append(match.group(1))
                departments.append(match.group(2))
                institutions.append(match.group(3))
                roles.append(match.group(4))
                if "chairman" in match.group(4).lower():
                    self.head= text.split(',')[0].strip()
                    print("Head Name: ",self.head)
        self.committee=names
        # Populate extracted data
        for i in range(len(names)):
            data = {
                "Name": names[i],
                "Title": titles[i],
                "Department": departments[i],
                "Institution": institutions[i],
                "Role": roles[i]
            }
            self.extracted_data.append(data)

    def create_excel(self):
        self.extract_information_from_docx()
        if not self.extracted_data:
            print("No data extracted. Run 'extract_information_from_docx' first.")
            return
        temp_folder = os.path.join(self.output_dir, "AllTables")  # Path to AllTables directory
        excel_path = os.path.join(temp_folder, "committee.xlsx")
        workbook = Workbook()
        sheet = workbook.active

        # Set column headers
        headers = ["Name", "Title", "Department", "Institution", "Role"]
        sheet.append(headers)

        # Add extracted data to the worksheet
        for item in self.extracted_data:
            row = [item["Name"], item["Title"], item["Department"], item["Institution"], item["Role"]]
            sheet.append(row)

        # Save the workbook as an Excel file
        workbook.save(excel_path)

    def extract_data_from_docx(self):
        # Function to extract data from a Word document
        try:
            self.extract_department_line()
            if self.dept:
                print("Department Line:", self.dept)
            else:
                print("No 'Department of' line found.")
        except FileNotFoundError:
            print(f"Error: The file '{self.docs_file}' was not found.")
        except Exception as e:
            print(f"Error: {e}")

        try:
            self.extract_year_and_term()
            if self.year and self.term:
                print("Word after 'bills :", self.year)
                print("Word after 'year':", self.term)
            else:
                print("No match found before the table.")
        except FileNotFoundError:
            print(f"Error: The file '{self.docx_file}' was not found.")
        except Exception as e:
            print(f"Error: {e}")

        
        print("Dept: ",self.dept)
        self.dept = self.dept_translate_to_bengali(self.dept.lower())
        print("Dept: ",self.dept)

        print("Year: ",self.year)
        self.year=self.convert_year_term_suffixes_to_bengali(self.year)
        print("Year: ",self.year)

        print("Term: ",self.term)
        self.term=self.convert_year_term_suffixes_to_bengali(self.term)
        print("Term: ",self.term)

        print("AD: ",self.AD)
        self.AD= "নিয়মিত পরীক্ষা " + bangla.convert_english_digit_to_bangla_digit(self.AD)
        print("AD: ",self.AD)

        # write at sample file
        # new_file_name = "_.xlsx"
        # Modify the code snippet where the Excel file is created and saved
        # temp_folder = os.path.join(self.output_dir, "AllTables")  # Path to AllTables directory
        # if not os.path.exists(temp_folder):
        #     os.makedirs(temp_folder)

        # Save the Excel file with some debug prints
        # shutil.copy(self.sample_excel, os.path.join(self.output_dir, new_file_name))
        # self.sample_excel=new_file_name
        file_path = os.path.join(self.output_dir, self.sample_excel)
        print(f"Debug: Saving Excel file at {file_path}")  # Add a debug print
        # wb = xw.Book()  # Create a new workbook
        wb = xw.Book(file_path)
        sheet = wb.sheets.active
        sheet.range('F3').value = self.AD
        sheet.range('G4').value = self.year
        sheet.range('I4').value = self.term
        sheet.range('F5').value ="বিভাগ: " + self.dept
        wb.save(file_path)
        wb.close()
        print(f"Debug: Excel file saved successfully at {file_path}")  # Add a debug print


        doc = Document(self.docx_file)
        text_content = ""
        self.tables_with_titles = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text)
                table_data.append(row_data)
            # Extracting the title before the table
            title = ""
            for paragraph in table.rows[0].cells[0].paragraphs:
                title += paragraph.text
            self.tables_with_titles.append({"Title": title, "Table": table_data})
            self.total_no_of_table = len(self.tables_with_titles)
        return text_content, self.tables_with_titles

    def generate_excel_from_docx(self, flag):
        if flag==0:
            self.name=""
        self.new_files = [] 
        self.progress_bar.pack(pady=20)
        if self.docx_file:
            self.output_dir = filedialog.askdirectory()
            if self.output_dir:
                text_content, self.tables_with_titles = self.extract_data_from_docx()
                temp_folder = os.path.join(self.output_dir, "AllTables")  # Path to AllTables directory
                if not os.path.exists(temp_folder):
                    os.makedirs(temp_folder)  # Create AllTables directory if it doesn't exist
                
                if self.tables_with_titles:
                    excel_path = os.path.join(temp_folder, "all_tables.xlsx")  # Save all_tables.xlsx inside AllTables
                    with pd.ExcelWriter(excel_path, engine='xlsxwriter') as writer:
                        for i, data in enumerate(self.tables_with_titles):
                            table = data["Table"]
                            df = pd.DataFrame(table)
                            df.ffill(axis=0, inplace=True)
                            sheet_name = f"Table_{i}"
                            df.to_excel(writer, sheet_name=sheet_name, index=False)
                            print(f"{sheet_name} added to Excel")
                        # messagebox.showinfo("Excel Created Successfully", f"All tables moved to {excel_path}!")
                        # self.clear_labels()
                    # self.pause_execution()
                else:
                    messagebox.showwarning("No Tables Found", "No tables were detected in the Word document.")
            else:
                messagebox.showwarning("Debug", "No output directory selected.")
        else:
            messagebox.showwarning("Oops!", "Please select a valid doc file.")
        self.create_excel()
        print("Committee table created....")
        self.process_first_table()

    def select_sample_excel(self):
        # Function to handle selection of Sample Excel file
        self.sample_excel = filedialog.askopenfilename(filetypes=[("Excel Files", "*.xlsx")])
        if self.sample_excel:
            self.sample_label.config(text=f"Selected Sample Excel: {self.sample_excel}")
            self.sample_label.pack()

    def select_docx(self):
        # Function to handle selection of Word document
        self.docx_file = filedialog.askopenfilename(filetypes=[("Word Files", "*.docx")])
        if self.docx_file:
            self.docx_label.config(text=f"Selected Word Doc: {self.docx_file}")
            self.docx_label.pack()


# <---------------------------------------------- definition section ----------------------------------------------> end
# <---------------------------------------------- class section ----------------------------------------------> end

# <---------------------------------------------- main function ----------------------------------------------> start
def main():
    root = tk.Tk()
    app = WordToExcelConverter(root)
    root.mainloop()
# <---------------------------------------------- main function ----------------------------------------------> end
if __name__ == "__main__":
    main()